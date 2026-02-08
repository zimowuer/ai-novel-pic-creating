# 导入基础库
import json          # 用于处理JSON数据（SD API返回结果解析）
import requests      # 用于发送HTTP请求（调用SD WebUI API）
import io            # 用于处理字节流（图片数据解码）
import base64        # 用于解码SD返回的base64格式图片
import threading     # 线程基础库（备用）
import random        # 用于随机选择图片宽高
import os            # 用于文件路径、目录操作
import shutil        # 用于复制文档文件
import time          # 用于重试机制的延时
from typing import List, Dict, Optional, Tuple  # 类型注解（提升代码可读性和健壮性）
from concurrent.futures import ThreadPoolExecutor  # 线程池（并发处理文本块）

# 导入第三方库
from PIL import Image               # 用于处理图片（保存SD生成的图片）
from docx import Document           # 用于读写docx文档（核心）
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 用于设置段落对齐方式（图片/文字居中）
from docx.shared import Inches      # 用于控制插入文档的图片尺寸
import tiktoken                     # OpenAI官方Token计算库（分割文本块）
from openai import OpenAI           # OpenAI Python客户端（调用文字API）
from openai import APIError, APITimeoutError  # OpenAI异常类（捕获API错误）

# ========== 通用重试装饰器（独立函数） ==========
def retry_decorator(retry_attr: str = "retry_times"):
    """
    【装饰器功能阐述】
    通用重试装饰器，专为类实例方法设计，实现API调用失败后的自动重试逻辑：
    - 从类实例中读取重试次数配置（默认读取retry_times属性）
    - 每次失败后延时1秒再重试
    - 重试次数用尽后抛出最终异常
    - 仅适配Doc2ImageGenerator类的实例方法
    
    :param retry_attr: 类实例中存储重试次数的属性名，默认"retry_times"
    """
    # 外层装饰器接收参数，返回内层装饰器
    def decorator(func):
        # 内层装饰器接收被装饰函数，返回包装函数
        def wrapper(*args, **kwargs):
            # 从被装饰函数的参数中获取类实例（第一个参数是self）
            self = args[0] if args else None
            # 校验装饰器使用场景（仅允许装饰类实例方法）
            if not self or not hasattr(self, retry_attr):
                raise ValueError("装饰器仅适用于Doc2ImageGenerator的实例方法")
            
            # 从类实例中读取重试次数配置
            retry_times = getattr(self, retry_attr)
            last_exception = None  # 记录最后一次异常
            
            # 重试循环：最多执行retry_times+1次（初始调用+retry_times次重试）
            for attempt in range(retry_times + 1):
                try:
                    # 执行原函数，返回结果（正常情况直接返回）
                    return func(*args, **kwargs)
                except Exception as e:
                    # 捕获异常，记录最后一次异常
                    last_exception = e
                    # 如果还有重试次数，打印提示并延时
                    if attempt < retry_times:
                        print(f"第{attempt+1}次调用失败，{e}，{retry_times - attempt}次重试机会...")
                        time.sleep(1)
                    else:
                        # 重试次数用尽，抛出最终异常（保留异常溯源）
                        raise Exception(f"重试{retry_times}次后仍失败：{e}") from last_exception
            # 理论上不会执行到这里，防止循环异常
            raise last_exception
        return wrapper
    return decorator

# ========== 核心类：文档转图片生成器 ==========
class Doc2ImageGenerator:
    """
    【类功能阐述】
    文档转图片生成器核心类，实现完整业务流程：
    1. 读取docx文档内容（排除空段落，支持表格内容）
    2. 按Token数分割文本为多个块（避免API Token超限）
    3. 调用OpenAI API生成SD提示词（全量传递角色提示词）
    4. 调用Stable Diffusion WebUI API生成图片
    5. 将图片直接插入docx文档，提示词保存到同级txt文件
    6. 支持并发处理、API超时重试、文档副本操作（避免修改原文档）
    
    核心修改点：
    - 角色提示词：取消正则匹配，改为全量拼接后传递给OpenAI API
    - 输出逻辑：图片插入docx，提示词保存到txt（文档同级目录）
    """
    def __init__(
        self,
        docx_path: str,
        token_per_chunk: int = 1000,
        openai_api_base: str = "https://api.openai.com/v1",
        openai_api_key: str = "",
        stable_api_url: str = "http://127.0.0.1:7860",
        sd_model_checkpoint: str = "",  # 唯一必填的SD参数
        concurrent_workers: int = 2,
        # 角色提示词字典（现在改为全量传递）
        character_prompts: Optional[Dict[str, str]] = None,
        negative_prompt: str = "ugly, blurry, low quality, deformed, disfigured",
        CLIP_stop_at_last_layers: int = 2,
        steps: int = 30,
        sampler_name: str = "Euler a",
        batch_size: int = 1,
        n_iter: int = 1,
        seed: int = 1,
        restore_faces: bool = False,
        width_height_list: List[tuple] = None,
        openai_timeout: float = 130.0,    # OpenAI API超时时间（默认130s）
        sd_timeout: float = 130.0,        # SD API超时时间（默认130s）
        retry_times: int = 2              # API调用失败重试次数（默认2次）
    ):
        """
        【初始化方法功能阐述】
        初始化生成器所有配置参数，完成参数校验和默认值设置：
        - 保存文档路径、Token分割数等基础配置
        - 初始化OpenAI/SD API的连接参数和生成参数
        - 校验必填参数（API密钥、SD模型名）
        - 设置默认图片宽高列表（确保≥512）
        - 初始化副本文档路径为None（后续生成）
        
        :param docx_path: 输入docx文档路径（必填）
        :param token_per_chunk: 每个文本块的Token数，默认1000
        :param openai_api_base: OpenAI API地址（本地代理可修改）
        :param openai_api_key: OpenAI API密钥（必填）
        :param stable_api_url: Stable WebUI API地址，默认本地http://127.0.0.1:7860
        :param sd_model_checkpoint: SD模型名称（如v1-5-pruned.ckpt，必填）
        :param concurrent_workers: 并发线程数，默认2（避免API过载）
        :param character_prompts: 角色-相貌提示词字典（全量传递）
        :param negative_prompt: 反向提示词（控制图片不生成的内容）
        :param CLIP_stop_at_last_layers: CLIP层数（SD参数）
        :param steps: 采样步数（SD参数，步数越多图片越精细）
        :param sampler_name: 采样器（SD参数，如Euler a、DPM++ 2M）
        :param batch_size: 批次大小（SD参数）
        :param n_iter: 迭代次数（SD参数）
        :param seed: 随机种子（SD参数，固定种子生成固定图片）
        :param restore_faces: 是否面部修复（SD参数）
        :param width_height_list: 图片宽高列表（默认4种尺寸）
        :param openai_timeout: OpenAI API超时时间（秒）
        :param sd_timeout: SD API超时时间（秒）
        :param retry_times: API调用失败重试次数
        """
        # 基础文档参数（保存到实例属性）
        self.docx_path = docx_path
        self.token_per_chunk = token_per_chunk
        
        # OpenAI API参数（保存到实例属性）
        self.openai_api_base = openai_api_base
        self.openai_api_key = openai_api_key
        # 角色提示词字典：如果传入None则初始化为空字典
        self.character_prompts = character_prompts if character_prompts is not None else {}
        self.openai_timeout = openai_timeout
        self.retry_times = retry_times
        
        # Stable WebUI API参数（保存到实例属性）
        self.stable_api_url = stable_api_url
        self.sd_model_checkpoint = sd_model_checkpoint
        self.negative_prompt = negative_prompt
        self.CLIP_stop_at_last_layers = CLIP_stop_at_last_layers
        self.steps = steps
        self.sampler_name = sampler_name
        self.batch_size = batch_size
        self.n_iter = n_iter
        self.seed = seed
        self.restore_faces = restore_faces
        self.sd_timeout = sd_timeout
        
        # 并发配置（线程池最大工作数）
        self.concurrent_workers = concurrent_workers
        
        # 宽高列表处理：设置默认值，校验最小尺寸≥512
        if width_height_list is None:
            self.width_height_list = [(512,512), (768,768), (896,896), (1024,1024)]
        else:
            self.width_height_list = []
            # 过滤并保留宽高≥512的尺寸（SD最小生成尺寸）
            for w, h in width_height_list:
                if w >= 512 and h >= 512:
                    self.width_height_list.append((w, h))
            # 如果过滤后为空，抛出异常（避免无效尺寸）
            if not self.width_height_list:
                raise ValueError("宽高列表中所有元素都过小，最小宽高需≥512")
        
        # 副本文档路径初始化（后续复制文档时赋值）
        self.docx_copy_path = None

        # 校验必填参数（防止运行时出错）
        if not self.openai_api_key:
            raise ValueError("OpenAI API密钥不能为空！")
        if not self.sd_model_checkpoint:
            raise ValueError("SD模型名称（sd_model_checkpoint）不能为空！")
        if not os.path.exists(self.docx_path):
            raise FileNotFoundError(f"文档不存在：{self.docx_path}")

    # ========== 私有方法：获取全量角色提示词 ==========
    def _get_all_character_prompts(self) -> str:
        """
        【方法功能阐述】
        核心修改点1：获取所有角色提示词并拼接成字符串，用于全量传递给OpenAI API：
        - 遍历character_prompts字典的所有值（角色对应的相貌提示词）
        - 用逗号+空格拼接所有提示词（符合SD提示词格式）
        - 空字典返回空字符串，避免传递无效内容
        - 打印拼接后的提示词，方便调试
        
        :return: 拼接后的全量角色提示词（空字典返回空字符串）
        """
        # 如果角色提示词字典为空，返回空字符串
        if not self.character_prompts:
            return ""
        # 拼接所有角色提示词的值（忽略键），用逗号分隔
        all_prompts = ", ".join(self.character_prompts.values())
        # 打印调试信息，方便查看传递的角色提示词
        print(f"全量角色提示词：{all_prompts}")
        return all_prompts

    # ========== 私有方法：复制原文档生成副本 ==========
    def _copy_docx_to_copy(self) -> str:
        """
        【方法功能阐述】
        复制原文档到同目录生成副本，避免修改原文档：
        - 解析原文档的目录、文件名、扩展名
        - 生成副本文件名：原文件名+_copy+扩展名
        - 用shutil.copy2复制文件（保留元数据）
        - 保存副本路径到实例属性，返回副本路径
        
        :return: 副本文档的完整路径
        """
        # 解析原文档的目录路径
        doc_dir = os.path.dirname(self.docx_path)
        # 解析原文档的文件名（含扩展名）
        doc_name = os.path.basename(self.docx_path)
        # 拆分文件名和扩展名
        doc_name_no_ext = os.path.splitext(doc_name)[0]
        doc_ext = os.path.splitext(doc_name)[1]
        
        # 生成副本文件名：原文件名_copy.扩展名
        copy_name = f"{doc_name_no_ext}_copy{doc_ext}"
        # 拼接副本的完整路径
        self.docx_copy_path = os.path.join(doc_dir, copy_name)
        
        # 复制原文档到副本路径（copy2保留文件元数据）
        shutil.copy2(self.docx_path, self.docx_copy_path)
        # 打印提示信息
        print(f"已复制原文档到副本：{self.docx_copy_path}")
        
        # 返回副本路径
        return self.docx_copy_path

    # ========== 私有方法：读取副本文档内容 ==========
    def _read_docx_content(self) -> Tuple[str, List[Tuple[int, str]]]:
        """
        【方法功能阐述】
        读取副本文档的正文内容（排除空段落、页眉页脚）：
        - 优先使用已生成的副本，未生成则先复制
        - 遍历文档所有段落，保留非空段落（记录段落索引和内容）
        - 遍历文档所有表格，保留非空单元格（记录单元格位置和内容）
        - 拼接所有内容为完整文本，校验非空
        - 返回完整文本和段落/表格内容列表
        
        :return: 元组(完整文本内容, 段落/表格内容列表)
        """
        # 如果副本路径未初始化，先复制文档生成副本
        if self.docx_copy_path is None:
            self._copy_docx_to_copy()
        # 打开副本文档
        doc = Document(self.docx_copy_path)
        
        # 存储所有非空内容（用于拼接完整文本）
        content = []
        # 存储段落索引和内容（格式：[(段落索引, 段落内容), ...]）
        paragraph_list = []
        
        # 遍历所有段落（排除空段落）
        for para_idx, para in enumerate(doc.paragraphs):
            # 去除首尾空格，判断是否为空
            para_text = para.text.strip()
            if para_text:
                # 添加到内容列表
                content.append(para_text)
                # 记录段落索引和内容
                paragraph_list.append((para_idx, para_text))
        
        # 遍历所有表格（处理表格中的文本）
        table_paragraphs = []
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # 去除首尾空格，判断是否为空
                    cell_text = cell.text.strip()
                    if cell_text:
                        # 添加到内容列表
                        content.append(cell_text)
                        # 记录表格单元格位置（自定义格式）和内容
                        table_paragraphs.append((f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}", cell_text))
        
        # 拼接所有内容为完整文本（换行分隔）
        full_content = "\n".join(content)
        # 校验内容非空（避免处理空文档）
        if not full_content:
            raise ValueError("文档正文为空，请检查文档内容！")
        
        # 返回完整文本 + 段落列表 + 表格段落列表
        return full_content, paragraph_list + table_paragraphs

    # ========== 私有方法：按Token分割文本块 ==========
    def _split_content_by_token(self, content: str, paragraph_list: List[Tuple[int, str]]) -> List[Dict]:
        """
        【方法功能阐述】
        按指定Token数分割文本为多个块，避免OpenAI API Token超限：
        - 使用tiktoken（OpenAI官方库）计算Token数（精准匹配API计费规则）
        - 按段落分割，避免拆分单个段落（保证语义完整）
        - 记录每个文本块的起始/结束段落索引、包含的段落列表
        - 返回分割后的文本块列表（每个块是字典，包含文本和位置信息）
        
        :param content: 完整文本内容（未使用，保留参数）
        :param paragraph_list: 段落/表格内容列表
        :return: 分割后的文本块列表（每个块是字典）
        """
        # 获取gpt-3.5-turbo的Token编码规则（精准计算Token数）
        encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
        
        # 存储分割后的文本块
        chunks = []
        # 存储当前块的Token列表（临时）
        current_chunk_tokens = []
        # 存储当前块的段落列表（临时）
        current_chunk_paragraphs = []
        # 存储当前块的Token总数（临时）
        current_token_count = 0
        
        # 遍历所有段落/表格单元格
        for para_idx, para_text in paragraph_list:
            # 计算当前段落的Token列表和数量
            para_tokens = encoding.encode(para_text)
            para_token_count = len(para_tokens)
            
            # 如果添加当前段落会超过Token上限，且当前块已有内容：保存当前块，重置临时变量
            if current_token_count + para_token_count > self.token_per_chunk and current_token_count > 0:
                # 拼接当前块的文本（换行分隔段落）
                chunk_text = "\n".join([p[1] for p in current_chunk_paragraphs])
                # 添加到文本块列表（记录文本、起始/结束索引、包含的段落）
                chunks.append({
                    "text": chunk_text,
                    "start_idx": current_chunk_paragraphs[0][0],
                    "end_idx": current_chunk_paragraphs[-1][0],
                    "paragraphs": current_chunk_paragraphs.copy()
                })
                # 重置临时变量，准备下一个块
                current_chunk_tokens = []
                current_chunk_paragraphs = []
                current_token_count = 0
            
            # 将当前段落添加到临时块中
            current_chunk_tokens.extend(para_tokens)
            current_chunk_paragraphs.append((para_idx, para_text))
            current_token_count += para_token_count
        
        # 处理最后一个块（循环结束后可能还有未保存的内容）
        if current_token_count > 0:
            chunk_text = "\n".join([p[1] for p in current_chunk_paragraphs])
            chunks.append({
                "text": chunk_text,
                "start_idx": current_chunk_paragraphs[0][0],
                "end_idx": current_chunk_paragraphs[-1][0],
                "paragraphs": current_chunk_paragraphs.copy()
            })
        
        # 返回分割后的文本块列表
        return chunks

    # ========== 私有方法：生成SD提示词（带重试） ==========
    @retry_decorator()
    def _generate_sd_prompt(self, chunk: str) -> str:
        """
        调用OpenAI API生成SD提示词（全量传递角色提示词，移除textarea标签）
        修复点：
        1. 提前初始化final_prompt，避免未赋值问题
        2. 补充API响应为空的校验
        3. 移除所有textarea标签相关逻辑
        4. 替换为新的基础提示词模板
        :param chunk: 单个文本块
        :return: 纯文本格式的SD提示词（无任何标签）
        """
        # 提前初始化变量，避免未赋值问题
        final_prompt = ""
        # 获取所有角色提示词（全量传递）
        all_character_prompt = self._get_all_character_prompts()
        
        # 替换为你提供的新基础提示词模板
        base_prompt_template = """
        请你协助完成Stable Diffusion文生图提示词生成任务，严格遵循以下引导和规则：
        
        1. 先理解文本切片：仔细阅读下方提供的文本切片内容，重点关注切片最接近末尾的描述部分——这是你需要生成提示词的核心依据；
        2. 场景选取要求：仅从切片最末尾的描述中，挑选1个具体、完整的场景（无需考虑前文内容，聚焦最后一个可视觉化的场景）；
        3. 提示词生成规则：
        - 必须用英文编写，以逗号分隔关键词/短语，仅保留表象化描述（如人物动作、服饰、环境、光影、物体形态等可直接视觉呈现的内容）；
        - 坚决拒绝包含情感、心理活动、抽象概念类词汇（如"happy"、"sad"、"brave"等）；
        - 提示词头部必须强制添加品质提升关键词：(masterpiece, best quality), beautiful detailed eyes, perfect face, detailed hair；
        - 你需要根据我所提供的所有样貌提示词结合你选择的画面中应该出现的人物，识别对应角色相貌提示词追加在提示词的适当位置来保证人物一致性，与其他关键词用逗号分隔；
        4. 输出要求：仅返回最终的提示词文本(请注意，提示词使用逗号分隔语言为英文)，无需任何额外解释、说明或格式修饰，确保可以直接用于Stable Diffusion生成图片。
        5.提示词使用逗号分隔语言为英文

        文本切片内容：
        {chunk_content}

        所有角色相貌提示词（追加到末尾）：{all_character_prompts}
        """
        
        # 填充模板（包含全量角色提示词）
        prompt_template = base_prompt_template.format(
            chunk_content=chunk,
            all_character_prompts=all_character_prompt
        )
        
        # 初始化OpenAI客户端
        client = OpenAI(
            api_key=self.openai_api_key,
            base_url=self.openai_api_base
        )
        
        try:
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "你是专业的Stable Diffusion提示词生成专家，擅长精准理解文本场景并转化为表象化英文提示词"},
                    {"role": "user", "content": prompt_template}
                ],
                temperature=0.7,
                max_tokens=1000,
                timeout=self.openai_timeout
            )
            
            # 补充响应为空的校验
            if not response or not response.choices or len(response.choices) == 0:
                raise Exception("OpenAI API返回空响应，未生成任何提示词")
            
            # 直接获取纯提示词（无textarea标签）
            raw_prompt = response.choices[0].message.content.strip()
            # 确保品质提升词在头部（二次校验）
            quality_prefix = "(masterpiece, best quality), beautiful detailed eyes, perfect face, detailed hair"
            if quality_prefix not in raw_prompt:
                raw_prompt = f"{quality_prefix}, {raw_prompt}"
            
            # 直接赋值为纯提示词，不再包裹任何标签
            final_prompt = raw_prompt
            
        except APITimeoutError:
            raise Exception(f"OpenAI API调用超时（超时时间：{self.openai_timeout}s）")
        except APIError as e:
            raise Exception(f"OpenAI API调用错误：{e}")
        except Exception as e:
            # 异常时赋值为错误提示（纯文本，无标签）
            final_prompt = f"生成提示词失败：{str(e)}"
            raise Exception(f"生成SD提示词失败：{str(e)}")
        
        # 最终兜底校验，确保返回值有效
        if not final_prompt:
            final_prompt = "未生成有效提示词"
        
        return final_prompt

    # ========== 私有方法：生成图片（带重试） ==========
    @retry_decorator()
    def _generate_image(self, prompt: str, chunk_index: int) -> str:
        """
        【方法功能阐述】
        调用Stable Diffusion WebUI API生成图片，保存到本地：
        - 随机选择图片宽高（从width_height_list中）
        - 构造SD API的请求参数（包含提示词、反向提示词、模型名等）
        - 发送POST请求调用txt2img接口（文生图）
        - 解码base64格式的图片数据，保存为PNG文件
        - 返回图片保存路径
        - 装饰器自动处理重试逻辑
        
        :param prompt: SD提示词（去除textarea标签后的纯文本）
        :param chunk_index: 文本块索引（用于生成图片文件名）
        :return: 图片保存的完整路径
        """
        # 随机选择图片宽高（从预设列表中）
        width, height = random.choice(self.width_height_list)
        
        # 构造SD API的请求参数（严格匹配SD WebUI的txt2img接口要求）
        payload = {
            "override_settings": {
                "sd_model_checkpoint": self.sd_model_checkpoint,  # 指定使用的SD模型
                "sd_vae": "animevae.pt",                          # VAE模型（提升图片色彩）
                "CLIP_stop_at_last_layers": self.CLIP_stop_at_last_layers,  # CLIP层数
            },
            "prompt": prompt,                  # 正向提示词
            "negative_prompt": self.negative_prompt,  # 反向提示词
            "steps": self.steps,               # 采样步数
            "sampler_name": self.sampler_name, # 采样器
            "width": width,                    # 图片宽度
            "height": height,                  # 图片高度
            "batch_size": self.batch_size,     # 批次大小
            "n_iter": self.n_iter,             # 迭代次数
            "seed": self.seed,                 # 随机种子
            "CLIP_stop_at_last_layers": self.CLIP_stop_at_last_layers,
            "restore_faces": self.restore_faces,  # 面部修复
        }
        
        try:
            # 发送POST请求调用SD WebUI的txt2img接口
            response = requests.post(
                url=f"{self.stable_api_url}/sdapi/v1/txt2img",  # API接口地址
                json=payload,                                   # 请求体（JSON格式）
                timeout=self.sd_timeout                         # 超时时间（秒）
            )
            # 校验响应状态码（非200则抛出异常）
            response.raise_for_status()
            
            # 解析JSON响应
            result = response.json()
            # 解码base64格式的图片数据（SD返回的第一个图片）
            image_data = base64.b64decode(result['images'][0])
            # 用PIL打开字节流图片
            image = Image.open(io.BytesIO(image_data))
            
            # 构造图片保存路径：文档目录/文档名_chunk_索引.png
            doc_dir = os.path.dirname(self.docx_path)
            doc_name = os.path.splitext(os.path.basename(self.docx_path))[0]
            image_path = os.path.join(doc_dir, f"{doc_name}_chunk_{chunk_index}.png")
            # 保存图片到本地
            image.save(image_path)
            
            # 返回图片路径
            return image_path
        
        # 捕获超时异常，抛出自定义提示
        except requests.exceptions.Timeout:
            raise Exception(f"SD API调用超时（超时时间：{self.sd_timeout}s）")
        # 捕获请求异常（如连接失败、状态码错误）
        except requests.exceptions.RequestException as e:
            raise Exception(f"SD API调用错误：{e}")
        # 捕获其他异常，抛出具体信息
        except Exception as e:
            raise Exception(f"生成图片失败（块{chunk_index}）：{str(e)}")

    # ========== 私有方法：写入文档+保存提示词 ==========
    def _write_to_docx(self, image_path: str, prompt: str, chunk: Dict, chunk_index: int):
        """
        【方法功能阐述】
        核心修改：仅在文档中插入居中的图片（移除所有引导词/提示词段落），提示词仍保存到txt文件：
        1. 提示词保存：追加到文档同级的txt文件（按文本块索引区分）
        2. 文档写入：仅在文本块后插入居中的图片，无任何引导文字
        3. 支持段落/表格两种文本块类型，保留原文档结构
        4. 保存修改后的文档，返回副本路径
        
        :param image_path: 图片保存路径
        :param prompt: 带textarea标签的提示词
        :param chunk: 文本块字典（包含起始/结束索引）
        :param chunk_index: 文本块索引
        :return: 修改后的副本文档路径
        """
        # 打开副本文档
        doc = Document(self.docx_copy_path)
        
        # ========== 保留提示词保存到txt文件的逻辑（不变） ==========
        doc_dir = os.path.dirname(self.docx_path)
        doc_name = os.path.splitext(os.path.basename(self.docx_path))[0]
        txt_path = os.path.join(doc_dir, f"{doc_name}_prompts.txt")
        with open(txt_path, "a", encoding="utf-8") as f:
            f.write(f"===== 文本块 {chunk_index} 提示词 =====\n")
            f.write(prompt.strip() + "\n\n")
        
        # 获取文本块的结束索引（用于确定插入位置）
        end_idx = chunk["end_idx"]
        
        # 处理普通段落文本块（索引为整数）
        if isinstance(end_idx, int):
            # 计算插入位置：文本块结束索引+1（在文本块后插入内容）
            insert_pos = end_idx + 1
            if insert_pos > len(doc.paragraphs):
                insert_pos = len(doc.paragraphs)
            
            # 仅创建空段落用于插入图片（移除所有引导文字）
            img_para = doc.add_paragraph()  # 空段落，无任何文字
            img_run = img_para.add_run()
            img_run.add_picture(image_path, width=Inches(6))  # 插入图片，宽度6英寸
            img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中
            
            # 段落移动函数（保留）
            def move_paragraph(para, new_index):
                doc._body._body.insert(new_index, para._element)
            
            # 仅移动图片段落到指定位置
            move_paragraph(img_para, insert_pos)
        
        # 处理表格文本块（索引为字符串）
        else:
            # 仅创建空段落插入图片（无任何引导文字）
            img_para = doc.add_paragraph()
            img_run = img_para.add_run()
            img_run.add_picture(image_path, width=Inches(6))
            img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 保存修改后的文档
        doc.save(self.docx_copy_path)
        
        # 打印提示信息
        print(f"文本块{chunk_index}：提示词已保存到{txt_path}，图片已插入文档（仅保留图片，无引导词）")
        return self.docx_copy_path

    # ========== 私有方法：处理单个文本块 ==========
    def _process_single_chunk(self, chunk: Dict, chunk_index: int):
        """
        处理单个文本块的完整流程：生成提示词→生成图片→插入文档+保存提示词
        修正点：移除textarea标签替换逻辑，直接使用纯提示词
        """
        try:
            print(f"开始处理文本块 {chunk_index}...")
            
            # 生成SD提示词（全量角色提示词，纯文本无标签）
            sd_prompt = self._generate_sd_prompt(chunk["text"])
            # 移除多余的标签替换步骤，直接使用纯提示词
            pure_prompt = sd_prompt
            
            # 生成图片
            image_path = self._generate_image(pure_prompt, chunk_index)
            
            # 插入图片+保存提示词
            processed_doc = self._write_to_docx(image_path, sd_prompt, chunk, chunk_index)
            
            print(f"文本块 {chunk_index} 处理完成！副本文档：{processed_doc}")
        
        except Exception as e:
            print(f"文本块 {chunk_index} 处理失败：{str(e)}")

    # ========== 主执行方法 ==========
    def run(self):
        """
        【方法功能阐述】
        生成器主执行方法，串联所有流程：
        1. 复制原文档生成副本（避免修改原文档）
        2. 读取副本文档内容（段落+表格）
        3. 按Token数分割文本为多个块
        4. 用线程池并发处理所有文本块（控制并发数）
        5. 等待所有并发任务完成，打印最终结果路径
        6. 捕获并抛出主流程异常
        
        执行流程：复制文档 → 读取内容 → 分割文本 → 并发处理 → 输出结果
        """
        try:
            # 第一步：复制原文档生成副本
            print("第一步：复制原文档生成副本...")
            self._copy_docx_to_copy()
            
            # 第二步：读取副本文档正文
            print("第二步：读取副本文档正文...")
            doc_content, paragraph_list = self._read_docx_content()
            
            # 第三步：分割文档为Token块
            print("第三步：分割文档为Token块...")
            text_chunks = self._split_content_by_token(doc_content, paragraph_list)
            print(f"文档分割完成，共生成 {len(text_chunks)} 个文本块")
            
            # 第四步：并发生成提示词和图片
            print("第四步：并发生成提示词和图片...")
            # 创建线程池（最大工作数=concurrent_workers）
            with ThreadPoolExecutor(max_workers=self.concurrent_workers) as executor:
                futures = []
                # 遍历所有文本块，提交到线程池
                for idx, chunk in enumerate(text_chunks):
                    future = executor.submit(self._process_single_chunk, chunk, idx)
                    futures.append(future)
                
                # 等待所有并发任务完成，捕获单个任务的异常
                for future in futures:
                    try:
                        future.result()
                    except Exception as e:
                        print(f"单个文本块处理失败：{str(e)}")
            
            # 打印最终结果路径
            print(f"所有处理任务已完成！")
            print(f"👉 最终文档：{self.docx_copy_path}")
            print(f"👉 提示词文件：{os.path.dirname(self.docx_path)}/{os.path.splitext(os.path.basename(self.docx_path))[0]}_prompts.txt")
        
        # 捕获主流程异常，打印并重新抛出
        except Exception as e:
            print(f"主流程执行失败：{str(e)}")
            raise